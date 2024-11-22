from markdown import markdown
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup  # To parse HTML content
import os
import re


def preprocess_custom_image_syntax(md_content):
    """
    Replace custom image syntax `![[image.png]]` with standard Markdown `![image](image.png)`
    """
    return re.sub(r'!\[\[(.+?)\]\]', r'![image](\1)', md_content)

def markdown_to_word(dir,md_file_name, output_docx_name, max_width_in_inches=6.0):
    md_file = os.path.join(dir,md_file_name)
    output_docx = os.path.join(dir,output_docx_name)
    # Step 1: Read the Markdown file
    with open(md_file, 'r', encoding='utf-8') as file:
        md_content = file.read()
    
       # Preprocess custom image syntax
    md_content = preprocess_custom_image_syntax(md_content)
    print(md_content)
    # Step 2: Convert Markdown to HTML
    html_content = markdown(md_content)
    print(html_content)
    # Step 3: Parse the HTML content for processing
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Step 4: Create a Word document
    doc = Document()
    
    for element in soup.contents:
        print("element: ", element.name)
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            # Process each child of the <p> tag in order
            for child in element.children:
                if child.name == 'img':
                    # Handle images
                    image_src = child['src']
                    image_path = os.path.join(dir, image_src)
                    if os.path.exists(image_path):
                        doc.add_picture(image_path, width=Inches(max_width_in_inches))
                    else:
                        doc.add_paragraph(f"[Image not found: {image_src}]")
                elif child.string:  # Add text content if present
                    doc.add_paragraph(child.string)
        elif element.name == 'ul':  # For unordered lists
            for li in element.find_all('li'):
                doc.add_paragraph(f"- {li.text}", style='List Bullet')
        elif element.name == 'ol':  # For ordered lists
            for li in element.find_all('li'):
                doc.add_paragraph(f"{li.text}", style='List Number')

    # Step 5: Save the Word document
    doc.save(output_docx)

# Example usage
main_dir_path = '/hri/storage/user/emenende/myNotes/dialogues/22_11_2024/dialogue1'

markdown_to_word(os.path.join(main_dir_path),'interaction_summary.md','output.docx')